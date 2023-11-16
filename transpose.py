import user_select as user
import music
from docx import Document
import re
import sys


def parse_chords_docx(file_path: str, chord_mapping: dict) -> list:
    doc = Document(file_path)
    converted_chords = {}
    for paragraph in doc.paragraphs:
        chord_matches = re.findall(r'\[([^][]+)]', paragraph.text)
        for item in chord_matches:
            item_cleaned, remainder = music.clean_chord(item)
            transposed_chord = chord_mapping[item_cleaned] + remainder
            #print(f'Original Symbol: {item}\t\tCleaned: {item_cleaned}\t\tRemainder {remainder}')
            converted_chords[item] = transposed_chord
    return converted_chords

def replace_chords_docx(file_path: str, chord_mapping: dict):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        for old_chord, new_chord in chord_mapping.items():
            for run in paragraph.runs:
                # Avoid replacing if the old chord is part of a longer chord
                if f'\[([{old_chord}^][]+)' in run.text and run.bold:
                    run.text = run.text.replace(f'[{old_chord}]', f'[{new_chord}]')
            
    doc.save("test.docx")

def debug_transpose(old_chords: list, delta_pitch_class, delta_semitones):
   for old_chord in old_chords:
    new_chord = music.transpose_chord(old_chord, delta_pitch_class, delta_semitones)
    print(f'Old Chord: {old_chord}       New Chord: {new_chord}')

def main(argv):
  valid_argv = [1, 2, 3]
  arg_len = len(argv)
  # Parse args
  if arg_len not in valid_argv:
    raise ValueError("Invalid number of command-line arguments. Please provide the required number of arguments.")     

  if arg_len==3:
    interval = music.intervals_array[int(argv[2])].value
    filepath = argv[1]
  elif arg_len==2:
    filepath = argv[1]
    interval = user.get_transpose_params()["Transpose"]
  else:
    filepath = user.open_file_dialog()
    interval = user.get_transpose_params()["Transpose"]

  # Pick transposition interval
       #consider edge cases here
  print(interval)
  delta_pitch_class = interval[1]
  delta_semitones = interval[0]

  # Create transposition dictionary
  chord_mapping = {}
  
  for old_chord in music.Note:
    new_chord = music.transpose_chord(old_chord.value[0], delta_pitch_class, delta_semitones)
    chord_mapping[old_chord.value[0]] = new_chord

  #print(transposed_chords)

  # For debug purposes only
  #debug_transpose(chords, delta_pitch_class, delta_semitones)

  # Parse document for chords
  chord_replacement_dictionary = parse_chords_docx(filepath, chord_mapping)  #may be obselete approach
  print(chord_replacement_dictionary)


  # Write new chords to document
  replace_chords_docx(filepath, chord_replacement_dictionary)



if __name__ == "__main__":
    main(sys.argv)