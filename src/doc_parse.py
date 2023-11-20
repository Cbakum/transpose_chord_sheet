import src.user_select as user
import src.music as music
from docx import Document
import re

def create_target_file_path(original_file_path: str) ->str:
    '''
    Creates the filename for the target chord sheet\n
    arg:
        original_file_path (str)- The name of the source .docx file\n
    returns:
        target_file_path (str)- The modified name to save the new .docx file
    '''

    # Find the index of "doc" in the file name
    index = original_file_path.find(".")

    if index != -1:
        # Extract the file name without the extension and text after "."
        file_name = original_file_path[:index]

        # Create the target file path by appending "_transposed.docx"
        target_file_path = f"{file_name}_transposed.docx"

    else:
        # Create dummy filepath to avoid overwriting source file
        target_file_path = "error_creating_filename.docx"
    
    return target_file_path

def parse_chords_docx(file_path: str, chord_mapping: dict):
    '''
    Looks through .docx file for [CHORDS] (only in square brackets),
    tranposes all chords, and rewrites tranposed chart to new file\n
    args:
            file_path (str)- The name of the source document\n
            chord_mapping (dict)- The translation dictionary from source note to target note
    '''
    doc = Document(file_path)
    converted_chords = {}
    for paragraph in doc.paragraphs:
        chord_matches = re.findall(r'\[([^][]+)]', paragraph.text)
        # Pass one to mark which chords have been replaced
        for item in chord_matches:
            # disconnect the root from quality in original chord
            root, quality, bass = music.clean_chord(item)
            # transpose root from mapping then reattach the quality
            transposed_chord = chord_mapping[root] + quality
            # If chord has inversion, transpose it and reattach
            if bass:
                transposed_chord = transposed_chord + chord_mapping[bass]
            # Add fully transposed chord to dictionary
            converted_chords[item] = transposed_chord

            # Attach * flag to indicate chord has been replaced
            # This avoids double replacing a chord on later loops, which would create incorrect transposition
            transposed_chord = f'{transposed_chord}*'

            # This replaces properly, but messes with formatting
            paragraph.text = paragraph.text.replace(f'[{item}]', f'[{transposed_chord}]')
            #print(f'{item} {transposed_chord}')

        # Pass two to remove * flag
        for item in chord_matches:
            # find the transposed chords again
            root, quality, bass = music.clean_chord(item)
            # transpose root from mapping then reattach the quality
            transposed_chord = chord_mapping[root] + quality
            # If chord has inversion, transpose it and reattach
            if bass:
                transposed_chord = transposed_chord + chord_mapping[bass]

            #remove the * flag that indicated chords were replaced
            paragraph.text = paragraph.text.replace(f'[{transposed_chord}*]', f'[{transposed_chord}]')
    # For debugging purposes
    #for paragraph in doc.paragraphs:
    #    for run in paragraph.runs:
    #        print(run.text)

    target_file_path = create_target_file_path(file_path)
    doc.save(target_file_path)


def reset_bold_chords(target_filepath: str):
    '''
    File write from parse_chords_docx does not persist bold text in chord symbols.
    This function parses through the document and restores boldness.\n
    arg:
        target_filepath (str)- the file to read from/ write to
    '''
    # To write later
    pass